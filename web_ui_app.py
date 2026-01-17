#!/usr/bin/env python3
"""
QA Documentation Generator - Web UI
Beautiful web interface for generating test plans and test cases
Built with Streamlit
"""

import streamlit as st
import os
import sys
import json
import PyPDF2
from anthropic import Anthropic
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import tempfile
from datetime import datetime
import io

# Page configuration
st.set_page_config(
    page_title="QA Docs Generator",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better UI
st.markdown("""
<style>
    .main-header {
        font-size: 3rem;
        font-weight: bold;
        text-align: center;
        color: #1f77b4;
        margin-bottom: 1rem;
    }
    .sub-header {
        font-size: 1.2rem;
        text-align: center;
        color: #666;
        margin-bottom: 2rem;
    }
    .stButton>button {
        width: 100%;
        background-color: #1f77b4;
        color: white;
        font-size: 1.1rem;
        padding: 0.75rem;
        border-radius: 8px;
    }
    .success-box {
        padding: 1rem;
        background-color: #d4edda;
        border-left: 5px solid #28a745;
        border-radius: 5px;
        margin: 1rem 0;
    }
    .info-box {
        padding: 1rem;
        background-color: #d1ecf1;
        border-left: 5px solid #17a2b8;
        border-radius: 5px;
        margin: 1rem 0;
    }
    .warning-box {
        padding: 1rem;
        background-color: #fff3cd;
        border-left: 5px solid #ffc107;
        border-radius: 5px;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Configuration
CLAUDE_API_KEY = os.getenv('ANTHROPIC_API_KEY')
MODEL_NAME = "claude-sonnet-4-20250514"


def read_pdf(pdf_file):
    """Extract text from uploaded PDF"""
    text = ""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        num_pages = len(pdf_reader.pages)
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            text += page.extract_text()
            progress = (page_num + 1) / num_pages
            progress_bar.progress(progress)
            status_text.text(f"📄 Reading page {page_num + 1}/{num_pages}")
        
        progress_bar.empty()
        status_text.empty()
        return text
    except Exception as e:
        st.error(f"❌ Error reading PDF: {e}")
        return None


def generate_test_plan_content(requirements_text, project_name):
    """Generate test plan using Claude API"""
    if not CLAUDE_API_KEY:
        st.error("❌ ANTHROPIC_API_KEY not set in environment variables!")
        return None
    
    client = Anthropic(api_key=CLAUDE_API_KEY)
    
    prompt = f"""You are a professional QA Test Plan writer. Based on the requirements document below, create a comprehensive test plan for the project "{project_name}".

REQUIREMENTS:
{requirements_text}

Generate a detailed test plan in JSON format with these sections (return ONLY JSON, no markdown):

{{
  "project_info": {{
    "name": "{project_name}",
    "version": "1.0",
    "prepared_by": "QA Team",
    "date": "{datetime.now().strftime('%Y-%m-%d')}",
    "test_environment": ["Development", "Pre-Production", "Production"]
  }},
  "description": "Detailed project description",
  "introduction": "Purpose and scope of testing",
  "goal": "Main testing objectives",
  "test_strategy": {{
    "functional": "Approach for functional testing",
    "integration": "Integration testing approach",
    "ui_ux": "UI/UX testing approach",
    "performance": "Performance testing approach",
    "security": "Security testing approach",
    "cross_platform": "Cross-platform testing approach",
    "regression": "Regression testing approach"
  }},
  "scope": {{
    "in_scope": ["List of in-scope items"],
    "out_of_scope": ["List of out-of-scope items"]
  }},
  "functional_requirements": [
    {{"req_id": "FR-001", "description": "Requirement description", "acceptance_criteria": "Criteria"}}
  ],
  "non_functional_requirements": [
    {{"category": "Performance", "requirement": "Details"}}
  ],
  "impact_zones": {{
    "red": ["High risk areas"],
    "yellow": ["Medium risk areas"],
    "green": ["Low risk areas"]
  }},
  "entry_criteria": ["List of entry criteria"],
  "exit_criteria": ["List of exit criteria"],
  "test_data": ["Test data requirements"],
  "testing_activities": [
    {{"phase": "Phase name", "activity": "Activity description", "timeline": "Duration"}}
  ],
  "roles": [
    {{"role": "Role name", "responsibility": "Responsibilities"}}
  ],
  "risks": [
    {{"risk": "Risk description", "mitigation": "Mitigation plan"}}
  ],
  "assumptions": ["List of assumptions"],
  "dependencies": ["List of dependencies"],
  "defect_management": {{
    "p1": "P1 definition and handling",
    "p2": "P2 definition and handling",
    "p3": "P3 definition and handling",
    "p4": "P4 definition and handling"
  }},
  "metrics": ["List of test metrics and KPIs"],
  "deliverables": ["List of test deliverables"],
  "limitations": ["Testing limitations"]
}}"""

    try:
        with st.spinner('🤖 Claude AI is generating test plan...'):
            response = client.messages.create(
                model=MODEL_NAME,
                max_tokens=16000,
                temperature=0.3,
                messages=[{"role": "user", "content": prompt}]
            )
        
        response_text = response.content[0].text
        
        # Clean response
        if response_text.strip().startswith('```'):
            response_text = response_text.strip()
            if response_text.startswith('```json'):
                response_text = response_text[7:]
            elif response_text.startswith('```'):
                response_text = response_text[3:]
            if response_text.endswith('```'):
                response_text = response_text[:-3]
            response_text = response_text.strip()
        
        test_plan = json.loads(response_text)
        return test_plan
    
    except Exception as e:
        st.error(f"❌ Error generating test plan: {e}")
        return None


def generate_test_cases_content(requirements_text, project_name):
    """Generate test cases using Claude API"""
    if not CLAUDE_API_KEY:
        st.error("❌ ANTHROPIC_API_KEY not set!")
        return None
    
    client = Anthropic(api_key=CLAUDE_API_KEY)
    
    prompt = f"""You are a professional QA Test Case writer. Based on the requirements, create comprehensive test cases for "{project_name}".

REQUIREMENTS:
{requirements_text}

Generate 40-60 detailed test cases in JSON format (return ONLY JSON array, no markdown):

[
  {{
    "id": "TC_001",
    "module": "Module/Feature Name",
    "title": "Test case title",
    "description": "What is being tested",
    "preconditions": "Pre-conditions (numbered list)",
    "steps": "Test execution steps (numbered)",
    "expected": "Expected results (numbered)",
    "priority": "P1 or P2 or P3",
    "type": "Functional/Integration/UI/Performance/Security/Cross-Platform/Edge Case/Regression",
    "platform": "Android/iOS/Both/Web"
  }}
]

Guidelines:
- Generate 40-60 test cases
- Cover all requirements
- Include positive, negative, edge cases
- Priority: P1 (Critical), P2 (High), P3 (Medium)
- Detailed numbered steps
- Type coverage: Functional, Integration, UI, Performance, Security"""

    try:
        with st.spinner('🤖 Claude AI is generating test cases...'):
            response = client.messages.create(
                model=MODEL_NAME,
                max_tokens=16000,
                temperature=0.3,
                messages=[{"role": "user", "content": prompt}]
            )
        
        response_text = response.content[0].text
        
        # Clean response
        if response_text.strip().startswith('```'):
            response_text = response_text.strip()
            if response_text.startswith('```json'):
                response_text = response_text[7:]
            elif response_text.startswith('```'):
                response_text = response_text[3:]
            if response_text.endswith('```'):
                response_text = response_text[:-3]
            response_text = response_text.strip()
        
        test_cases = json.loads(response_text)
        return test_cases
    
    except Exception as e:
        st.error(f"❌ Error generating test cases: {e}")
        return None


def create_word_document(test_plan, project_name):
    """Create Word document from test plan"""
    doc = Document()
    
    # Title
    title = doc.add_heading(f'{project_name} - QA Test Plan', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Project Info Table
    doc.add_heading('1. Project Information', level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    info = test_plan.get('project_info', {})
    table.rows[0].cells[0].text = 'Project Name'
    table.rows[0].cells[1].text = info.get('name', project_name)
    table.rows[1].cells[0].text = 'Version'
    table.rows[1].cells[1].text = info.get('version', '1.0')
    table.rows[2].cells[0].text = 'Prepared By'
    table.rows[2].cells[1].text = info.get('prepared_by', 'QA Team')
    table.rows[3].cells[0].text = 'Date'
    table.rows[3].cells[1].text = info.get('date', datetime.now().strftime('%Y-%m-%d'))
    table.rows[4].cells[0].text = 'Test Environment'
    table.rows[4].cells[1].text = ', '.join(info.get('test_environment', []))
    
    # Description
    doc.add_heading('2. Description', level=1)
    doc.add_paragraph(test_plan.get('description', 'N/A'))
    
    # Introduction
    doc.add_heading('3. Introduction', level=1)
    doc.add_paragraph(test_plan.get('introduction', 'N/A'))
    
    # Goal
    doc.add_heading('4. Goal', level=1)
    doc.add_paragraph(test_plan.get('goal', 'N/A'))
    
    # Test Strategy
    doc.add_heading('5. Test Strategy', level=1)
    strategy = test_plan.get('test_strategy', {})
    if strategy:
        doc.add_heading('5.1 Functional Testing', level=2)
        doc.add_paragraph(strategy.get('functional', 'N/A'))
        doc.add_heading('5.2 Integration Testing', level=2)
        doc.add_paragraph(strategy.get('integration', 'N/A'))
        doc.add_heading('5.3 UI/UX Testing', level=2)
        doc.add_paragraph(strategy.get('ui_ux', 'N/A'))
        doc.add_heading('5.4 Performance Testing', level=2)
        doc.add_paragraph(strategy.get('performance', 'N/A'))
        doc.add_heading('5.5 Security Testing', level=2)
        doc.add_paragraph(strategy.get('security', 'N/A'))
        doc.add_heading('5.6 Cross-Platform Testing', level=2)
        doc.add_paragraph(strategy.get('cross_platform', 'N/A'))
        doc.add_heading('5.7 Regression Testing', level=2)
        doc.add_paragraph(strategy.get('regression', 'N/A'))
    
    # Test Scope
    doc.add_heading('6. Test Scope', level=1)
    scope = test_plan.get('scope', {})
    if scope:
        doc.add_heading('6.1 In-Scope', level=2)
        for item in scope.get('in_scope', []):
            doc.add_paragraph(item, style='List Bullet')
        doc.add_heading('6.2 Out-of-Scope', level=2)
        for item in scope.get('out_of_scope', []):
            doc.add_paragraph(item, style='List Bullet')
    
    # Functional Requirements
    doc.add_heading('7. Functional Requirements', level=1)
    func_reqs = test_plan.get('functional_requirements', [])
    if func_reqs:
        for req in func_reqs:
            doc.add_heading(f"{req.get('req_id', 'FR-XXX')}: {req.get('description', 'N/A')[:50]}...", level=2)
            doc.add_paragraph(f"Description: {req.get('description', 'N/A')}")
            doc.add_paragraph(f"Acceptance Criteria: {req.get('acceptance_criteria', 'N/A')}")
    
    # Non-Functional Requirements
    doc.add_heading('8. Non-Functional Requirements', level=1)
    non_func_reqs = test_plan.get('non_functional_requirements', [])
    if non_func_reqs:
        for req in non_func_reqs:
            doc.add_paragraph(f"{req.get('category', 'N/A')}: {req.get('requirement', 'N/A')}", style='List Bullet')
    
    # Impact Zones
    doc.add_heading('9. Impact Zones', level=1)
    impact = test_plan.get('impact_zones', {})
    if impact:
        doc.add_heading('9.1 Red Zone (High Risk)', level=2)
        for item in impact.get('red', []):
            doc.add_paragraph(item, style='List Bullet')
        doc.add_heading('9.2 Yellow Zone (Medium Risk)', level=2)
        for item in impact.get('yellow', []):
            doc.add_paragraph(item, style='List Bullet')
        doc.add_heading('9.3 Green Zone (Low Risk)', level=2)
        for item in impact.get('green', []):
            doc.add_paragraph(item, style='List Bullet')
    
    # Entry Criteria
    doc.add_heading('10. Entry Criteria', level=1)
    for item in test_plan.get('entry_criteria', []):
        doc.add_paragraph(item, style='List Bullet')
    
    # Exit Criteria
    doc.add_heading('11. Exit Criteria', level=1)
    for item in test_plan.get('exit_criteria', []):
        doc.add_paragraph(item, style='List Bullet')
    
    # Test Data Requirements
    doc.add_heading('12. Test Data Requirements', level=1)
    for item in test_plan.get('test_data', []):
        doc.add_paragraph(item, style='List Bullet')
    
    # Test Environment
    doc.add_heading('13. Test Environment', level=1)
    doc.add_paragraph(', '.join(info.get('test_environment', [])))
    
    # Testing Activities
    doc.add_heading('14. Testing Activities', level=1)
    activities = test_plan.get('testing_activities', [])
    if activities:
        table = doc.add_table(rows=len(activities)+1, cols=3)
        table.style = 'Light Grid Accent 1'
        table.rows[0].cells[0].text = 'Phase'
        table.rows[0].cells[1].text = 'Activity'
        table.rows[0].cells[2].text = 'Timeline'
        for i, activity in enumerate(activities, 1):
            table.rows[i].cells[0].text = activity.get('phase', 'N/A')
            table.rows[i].cells[1].text = activity.get('activity', 'N/A')
            table.rows[i].cells[2].text = activity.get('timeline', 'N/A')
    
    # Roles and Responsibilities
    doc.add_heading('15. Roles and Responsibilities', level=1)
    roles = test_plan.get('roles', [])
    if roles:
        for role in roles:
            doc.add_heading(role.get('role', 'N/A'), level=2)
            doc.add_paragraph(role.get('responsibility', 'N/A'))
    
    # Risks
    doc.add_heading('16. Risks and Mitigation', level=1)
    risks = test_plan.get('risks', [])
    if risks:
        for risk in risks:
            doc.add_paragraph(f"Risk: {risk.get('risk', 'N/A')}", style='List Bullet')
            doc.add_paragraph(f"Mitigation: {risk.get('mitigation', 'N/A')}", style='List Bullet 2')
    
    # Assumptions
    doc.add_heading('17. Assumptions', level=1)
    for item in test_plan.get('assumptions', []):
        doc.add_paragraph(item, style='List Bullet')
    
    # Dependencies
    doc.add_heading('18. Dependencies', level=1)
    for item in test_plan.get('dependencies', []):
        doc.add_paragraph(item, style='List Bullet')
    
    # Defect Management
    doc.add_heading('19. Defect Management', level=1)
    defect = test_plan.get('defect_management', {})
    if defect:
        doc.add_heading('P1 - Critical', level=2)
        doc.add_paragraph(defect.get('p1', 'N/A'))
        doc.add_heading('P2 - High', level=2)
        doc.add_paragraph(defect.get('p2', 'N/A'))
        doc.add_heading('P3 - Medium', level=2)
        doc.add_paragraph(defect.get('p3', 'N/A'))
        doc.add_heading('P4 - Low', level=2)
        doc.add_paragraph(defect.get('p4', 'N/A'))
    
    # Test Metrics
    doc.add_heading('20. Test Metrics and KPIs', level=1)
    for item in test_plan.get('metrics', []):
        doc.add_paragraph(item, style='List Bullet')
    
    # Deliverables
    doc.add_heading('21. Deliverables', level=1)
    for item in test_plan.get('deliverables', []):
        doc.add_paragraph(item, style='List Bullet')
    
    # Limitations
    doc.add_heading('22. Limitations and Exclusions', level=1)
    for item in test_plan.get('limitations', []):
        doc.add_paragraph(item, style='List Bullet')
    
    # Approval Section
    doc.add_heading('23. Approval', level=1)
    approval_table = doc.add_table(rows=4, cols=3)
    approval_table.style = 'Light Grid Accent 1'
    approval_table.rows[0].cells[0].text = 'Role'
    approval_table.rows[0].cells[1].text = 'Name'
    approval_table.rows[0].cells[2].text = 'Signature & Date'
    approval_table.rows[1].cells[0].text = 'QA Lead'
    approval_table.rows[2].cells[0].text = 'Project Manager'
    approval_table.rows[3].cells[0].text = 'Product Owner'
    
    # Convert to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes


def create_excel_file(test_cases, project_name):
    """Create Excel file from test cases"""
    wb = Workbook()
    sheet = wb.active
    sheet.title = "Test Cases"
    
    # Headers
    headers = ['Test Case ID', 'Module', 'Test Case Title', 'Description', 
               'Pre-conditions', 'Test Steps', 'Expected Results', 
               'Priority', 'Test Type', 'Platform']
    
    header_font = Font(bold=True, color='FFFFFF', size=11)
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    
    for col, header in enumerate(headers, 1):
        cell = sheet.cell(row=1, column=col)
        cell.value = header
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    
    # Data
    for idx, tc in enumerate(test_cases, start=2):
        sheet[f'A{idx}'] = tc.get('id', f'TC_{idx-1:03d}')
        sheet[f'B{idx}'] = tc.get('module', '')
        sheet[f'C{idx}'] = tc.get('title', '')
        sheet[f'D{idx}'] = tc.get('description', '')
        sheet[f'E{idx}'] = tc.get('preconditions', '')
        sheet[f'F{idx}'] = tc.get('steps', '')
        sheet[f'G{idx}'] = tc.get('expected', '')
        sheet[f'H{idx}'] = tc.get('priority', 'P2')
        sheet[f'I{idx}'] = tc.get('type', 'Functional')
        sheet[f'J{idx}'] = tc.get('platform', 'Both')
    
    # Convert to bytes
    excel_bytes = io.BytesIO()
    wb.save(excel_bytes)
    excel_bytes.seek(0)
    return excel_bytes


def main():
    # Initialize session state for generated files
    if 'test_plan_docx' not in st.session_state:
        st.session_state.test_plan_docx = None
    if 'test_plan_json' not in st.session_state:
        st.session_state.test_plan_json = None
    if 'test_cases_xlsx' not in st.session_state:
        st.session_state.test_cases_xlsx = None
    if 'test_cases_json' not in st.session_state:
        st.session_state.test_cases_json = None
    if 'project_name' not in st.session_state:
        st.session_state.project_name = None
    if 'test_cases_count' not in st.session_state:
        st.session_state.test_cases_count = 0
    if 'test_cases_stats' not in st.session_state:
        st.session_state.test_cases_stats = {}
    
    # Header
    st.markdown('<p class="main-header">🚀 QA Documentation Generator</p>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">Automatically generate Test Plans and Test Cases using AI</p>', unsafe_allow_html=True)
    
    # Sidebar
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # API Key check
        if CLAUDE_API_KEY:
            st.success("✅ API Key Configured")
        else:
            st.error("❌ API Key Not Set")
            st.info("Set ANTHROPIC_API_KEY environment variable")
        
        st.markdown("---")
        st.header("📚 About")
        st.markdown("""
        This tool uses **Claude AI** to generate:
        - 📄 Professional Test Plans
        - 🧪 Comprehensive Test Cases
        
        **Powered by:**
        - Anthropic Claude API
        - Streamlit
        """)
        
        st.markdown("---")
        st.header("💡 Tips")
        st.markdown("""
        - Upload clear PDF requirements
        - Use descriptive project names
        - Review generated docs
        - Customize as needed
        """)
    
    # Main content
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("📤 Upload Requirements")
        uploaded_file = st.file_uploader(
            "Upload your requirements PDF",
            type=['pdf'],
            help="Upload a PDF file containing project requirements"
        )
    
    with col2:
        st.header("📝 Project Name")
        project_name = st.text_input(
            "Enter project name",
            value="My Project",
            help="Enter a descriptive name for your project"
        )
    
    if uploaded_file:
        st.markdown('<div class="info-box">✅ PDF Uploaded Successfully</div>', unsafe_allow_html=True)
        
        # What to generate
        st.header("🎯 What would you like to generate?")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            generate_plan = st.checkbox("📄 Test Plan", value=True)
        with col2:
            generate_cases = st.checkbox("🧪 Test Cases", value=True)
        with col3:
            st.write("")  # Spacer
        
        if not (generate_plan or generate_cases):
            st.warning("⚠️ Please select at least one option")
            return
        
        # Generate button
        st.markdown("---")
        if st.button("🚀 Generate Documentation", type="primary"):
            if not CLAUDE_API_KEY:
                st.error("❌ Please set ANTHROPIC_API_KEY environment variable first!")
                return
            
            # Clear previous files
            st.session_state.test_plan_docx = None
            st.session_state.test_plan_json = None
            st.session_state.test_cases_xlsx = None
            st.session_state.test_cases_json = None
            st.session_state.project_name = project_name
            st.session_state.test_cases_count = 0
            st.session_state.test_cases_stats = {}
            
            # Read PDF
            st.header("📖 Processing Requirements")
            pdf_text = read_pdf(uploaded_file)
            
            if not pdf_text:
                return
            
            st.success(f"✅ Extracted {len(pdf_text)} characters from PDF")
            
            # Generate Test Plan
            if generate_plan:
                st.header("📄 Generating Test Plan")
                test_plan = generate_test_plan_content(pdf_text, project_name)
                
                if test_plan:
                    st.success("✅ Test Plan generated successfully!")
                    
                    # Create Word document
                    with st.spinner("📝 Creating Word document..."):
                        doc_bytes = create_word_document(test_plan, project_name)
                    
                    # Store in session state
                    st.session_state.test_plan_docx = doc_bytes.getvalue()
                    
                    # Create JSON
                    json_bytes = json.dumps(test_plan, indent=2, ensure_ascii=False).encode('utf-8')
                    st.session_state.test_plan_json = json_bytes
            
            # Generate Test Cases
            if generate_cases:
                st.header("🧪 Generating Test Cases")
                test_cases = generate_test_cases_content(pdf_text, project_name)
                
                if test_cases:
                    st.success(f"✅ Generated {len(test_cases)} test cases!")
                    
                    # Store count and stats
                    st.session_state.test_cases_count = len(test_cases)
                    
                    # Calculate statistics
                    priorities = {}
                    types = {}
                    
                    for tc in test_cases:
                        p = tc.get('priority', 'P2')
                        priorities[p] = priorities.get(p, 0) + 1
                        t = tc.get('type', 'Functional')
                        types[t] = types.get(t, 0) + 1
                    
                    st.session_state.test_cases_stats = {
                        'priorities': priorities,
                        'types': types
                    }
                    
                    # Create Excel
                    with st.spinner("📊 Creating Excel file..."):
                        excel_bytes = create_excel_file(test_cases, project_name)
                    
                    # Store in session state
                    st.session_state.test_cases_xlsx = excel_bytes.getvalue()
                    
                    # Create JSON
                    json_bytes = json.dumps(test_cases, indent=2, ensure_ascii=False).encode('utf-8')
                    st.session_state.test_cases_json = json_bytes
            
            # Success message
            st.markdown('<div class="success-box">🎉 <strong>Generation Complete!</strong><br>Scroll down to download your files.</div>', unsafe_allow_html=True)
            st.rerun()  # Rerun to show download section
        
        # Display download buttons if files exist in session state
        if st.session_state.test_plan_docx or st.session_state.test_cases_xlsx:
            st.markdown("---")
            st.header("📥 Download Your Files")
            
            # Test Plan Downloads
            if st.session_state.test_plan_docx:
                st.subheader("📄 Test Plan")
                col1, col2 = st.columns(2)
                
                with col1:
                    st.download_button(
                        label="📥 Download Test Plan (Word)",
                        data=st.session_state.test_plan_docx,
                        file_name=f"{st.session_state.project_name.replace(' ', '_')}_Test_Plan.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_plan_docx"
                    )
                
                with col2:
                    st.download_button(
                        label="📥 Download Test Plan (JSON)",
                        data=st.session_state.test_plan_json,
                        file_name=f"{st.session_state.project_name.replace(' ', '_')}_Test_Plan.json",
                        mime="application/json",
                        key="download_plan_json"
                    )
            
            # Test Cases Downloads
            if st.session_state.test_cases_xlsx:
                st.subheader("🧪 Test Cases")
                
                # Show statistics
                if st.session_state.test_cases_stats:
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        st.metric("Total Test Cases", st.session_state.test_cases_count)
                    with col2:
                        priorities = st.session_state.test_cases_stats.get('priorities', {})
                        st.metric("P1 (Critical)", priorities.get('P1', 0))
                    with col3:
                        types = st.session_state.test_cases_stats.get('types', {})
                        st.metric("Test Types", len(types))
                
                # Download buttons
                col1, col2 = st.columns(2)
                
                with col1:
                    st.download_button(
                        label="📥 Download Test Cases (Excel)",
                        data=st.session_state.test_cases_xlsx,
                        file_name=f"{st.session_state.project_name.replace(' ', '_')}_Test_Cases.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="download_cases_xlsx"
                    )
                
                with col2:
                    st.download_button(
                        label="📥 Download Test Cases (JSON)",
                        data=st.session_state.test_cases_json,
                        file_name=f"{st.session_state.project_name.replace(' ', '_')}_Test_Cases.json",
                        mime="application/json",
                        key="download_cases_json"
                    )
            
            # Clear button
            st.markdown("---")
            if st.button("🔄 Generate New Documents", type="secondary"):
                st.session_state.test_plan_docx = None
                st.session_state.test_plan_json = None
                st.session_state.test_cases_xlsx = None
                st.session_state.test_cases_json = None
                st.session_state.project_name = None
                st.session_state.test_cases_count = 0
                st.session_state.test_cases_stats = {}
                st.rerun()
    
    else:
        # Instructions
        st.info("""
        ### 👋 Welcome to QA Documentation Generator!
        
        **How to use:**
        1. 📤 Upload your requirements PDF
        2. ✏️ Enter a project name
        3. ✅ Select what to generate (Test Plan / Test Cases)
        4. 🚀 Click "Generate Documentation"
        5. 📥 Download your generated files
        
        **Requirements:**
        - Set `ANTHROPIC_API_KEY` environment variable
        - Upload PDF with clear requirements
        - Good internet connection
        """)


if __name__ == "__main__":
    main()

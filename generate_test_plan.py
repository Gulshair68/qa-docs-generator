#!/usr/bin/env python3
"""
Automated Test Plan Generator using Claude API
Author: Created for QA Team
Description: Takes requirement PDF, generates comprehensive test plan using Claude API
"""

import os
import sys
import json
import PyPDF2
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configuration
CLAUDE_API_KEY = os.getenv('ANTHROPIC_API_KEY')  # Set your API key as environment variable
MODEL_NAME = "claude-sonnet-4-20250514"  # Latest Sonnet model


def read_pdf(pdf_path):
    """Extract text from PDF file"""
    print(f"📄 Reading PDF: {pdf_path}")
    text = ""
    
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            print(f"   Total pages: {num_pages}")
            
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
                print(f"   Extracted page {page_num + 1}/{num_pages}")
        
        print(f"✅ Successfully extracted {len(text)} characters")
        return text
    
    except Exception as e:
        print(f"❌ Error reading PDF: {e}")
        sys.exit(1)


def generate_test_plan_with_claude(requirements_text, project_name="Project"):
    """Generate test plan using Claude API"""
    print(f"\n🤖 Calling Claude API to generate test plan...")
    
    if not CLAUDE_API_KEY:
        print("❌ Error: ANTHROPIC_API_KEY environment variable not set!")
        print("   Set it using: export ANTHROPIC_API_KEY='your-api-key'")
        sys.exit(1)
    
    client = Anthropic(api_key=CLAUDE_API_KEY)
    
    prompt = f"""You are a professional QA Test Plan writer. Based on the following requirements document, create a comprehensive QA Test Plan.

REQUIREMENTS DOCUMENT:
{requirements_text}

Please generate a detailed test plan with the following sections in JSON format:

{{
  "project_name": "Extract from requirements or use '{project_name}'",
  "version": "1.0",
  "description": "Brief description of the project and what is being tested",
  "introduction": "Introduction explaining the purpose of testing",
  "goal": "Clear testing objectives",
  "test_strategy": [
    "List of testing types: Functional, Integration, UI/UX, Performance, Security, Cross-Platform, Regression, etc."
  ],
  "in_scope": [
    "List of features/requirements that are in scope for testing"
  ],
  "out_of_scope": [
    "List of items that are out of scope"
  ],
  "functional_requirements": [
    {{
      "id": "Requirement ID",
      "title": "Requirement title",
      "description": "Detailed description",
      "acceptance_criteria": ["Criteria 1", "Criteria 2"]
    }}
  ],
  "non_functional_requirements": [
    "Performance: Load time < X seconds",
    "Reliability: Success rate >= Y%",
    "etc."
  ],
  "impact_zones": {{
    "red": ["Critical areas that must be tested thoroughly"],
    "yellow": ["Medium impact - sanity testing required"],
    "green": ["Low impact - smoke testing sufficient"]
  }},
  "entry_criteria": [
    "Criteria that must be met before testing starts"
  ],
  "exit_criteria": [
    "Criteria that must be met to complete testing"
  ],
  "test_data_requirements": [
    "Required test data"
  ],
  "test_environment": [
    {{"name": "Dev", "purpose": "Development testing"}},
    {{"name": "Pre-Prod", "purpose": "UAT"}},
    {{"name": "Production", "purpose": "Sanity"}}
  ],
  "testing_activities": [
    {{
      "activity": "Activity name",
      "details": "Details",
      "duration": "X hours/days"
    }}
  ],
  "roles_responsibilities": [
    {{"role": "QA Lead", "name": "TBD", "responsibilities": "Test strategy, sign-off"}},
    {{"role": "QA Engineer", "name": "Names", "responsibilities": "Test execution"}}
  ],
  "risks": [
    "Risk 1: Description",
    "Risk 2: Description"
  ],
  "assumptions": [
    "Assumption 1",
    "Assumption 2"
  ],
  "dependencies": [
    "Dependency 1",
    "Dependency 2"
  ],
  "defect_management": [
    "P1: Critical - description",
    "P2: High - description",
    "P3: Medium - description",
    "P4: Low - description"
  ],
  "test_metrics": [
    "Test coverage: Target",
    "Pass rate: Target"
  ],
  "deliverables": [
    "Daily reports",
    "Test execution report",
    "etc."
  ],
  "limitations": [
    "Limitation 1",
    "Limitation 2"
  ]
}}

Please analyze the requirements thoroughly and provide a comprehensive test plan in valid JSON format only. Do not include any markdown formatting or code blocks - just pure JSON."""

    try:
        response = client.messages.create(
            model=MODEL_NAME,
            max_tokens=16000,
            temperature=0.3,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        
        response_text = response.content[0].text
        print("✅ Received response from Claude")
        
        # Clean response if it has markdown code blocks
        if response_text.strip().startswith('```'):
            # Remove ```json and ``` markers
            response_text = response_text.strip()
            if response_text.startswith('```json'):
                response_text = response_text[7:]
            elif response_text.startswith('```'):
                response_text = response_text[3:]
            if response_text.endswith('```'):
                response_text = response_text[:-3]
            response_text = response_text.strip()
        
        # Parse JSON
        test_plan = json.loads(response_text)
        print("✅ Successfully parsed test plan JSON")
        return test_plan
    
    except json.JSONDecodeError as e:
        print(f"❌ Error parsing JSON: {e}")
        print(f"Response text: {response_text[:500]}...")
        sys.exit(1)
    
    except Exception as e:
        print(f"❌ Error calling Claude API: {e}")
        sys.exit(1)


def set_cell_border(cell, **kwargs):
    """Set cell border"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create borders element
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), 'CCCCCC')
            tcBorders.append(edge_el)
    
    tcPr.append(tcBorders)


def create_word_document(test_plan, output_path):
    """Create professional Word document from test plan"""
    print(f"\n📝 Creating Word document...")
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading(f"{test_plan.get('project_name', 'Project')} - QA Test Plan", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph(f"Version {test_plan.get('version', '1.0')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()
    
    # Project Information Table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    info_data = [
        ('Project Name', test_plan.get('project_name', 'N/A')),
        ('Document Status', 'DRAFT'),
        ('Version', test_plan.get('version', '1.0')),
        ('Team Members', 'QA Team'),
        ('Test Environment', ', '.join([env['name'] for env in test_plan.get('test_environment', [])]))
    ]
    
    for i, (label, value) in enumerate(info_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[i].cells[1].text = str(value)
    
    doc.add_paragraph()
    
    # Description Section
    doc.add_heading('Description', level=1)
    doc.add_paragraph(test_plan.get('description', 'N/A'))
    
    # Introduction Section
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(test_plan.get('introduction', 'N/A'))
    
    # Goal Section
    doc.add_heading('Goal', level=1)
    doc.add_paragraph(test_plan.get('goal', 'N/A'))
    
    # Test Strategy
    doc.add_heading('Test Strategy', level=1)
    for strategy in test_plan.get('test_strategy', []):
        p = doc.add_paragraph(strategy, style='List Bullet')
    
    # Test Scope
    doc.add_heading('Test Scope', level=1)
    
    doc.add_heading('In-Scope', level=2)
    for item in test_plan.get('in_scope', []):
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Out-of-Scope', level=2)
    for item in test_plan.get('out_of_scope', []):
        doc.add_paragraph(item, style='List Bullet')
    
    # Functional Requirements
    doc.add_heading('Functional Requirements', level=1)
    for idx, req in enumerate(test_plan.get('functional_requirements', []), 1):
        doc.add_heading(f"{idx}. {req.get('title', 'Requirement')}", level=2)
        doc.add_paragraph(req.get('description', ''))
        
        if req.get('acceptance_criteria'):
            doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
            for criteria in req['acceptance_criteria']:
                p = doc.add_paragraph(criteria, style='List Bullet 2')
    
    # Non-Functional Requirements
    doc.add_heading('Non-Functional Requirements', level=1)
    for nfr in test_plan.get('non_functional_requirements', []):
        doc.add_paragraph(nfr, style='List Bullet')
    
    # Impact Zones
    doc.add_heading('Impacted Areas', level=1)
    
    impact = test_plan.get('impact_zones', {})
    
    doc.add_heading('Red Zones (Critical)', level=2)
    for item in impact.get('red', []):
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Yellow Zones (Medium Impact)', level=2)
    for item in impact.get('yellow', []):
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Green Zones (Low Impact)', level=2)
    for item in impact.get('green', []):
        doc.add_paragraph(item, style='List Bullet')
    
    # Entry & Exit Criteria
    doc.add_heading('Entry & Exit Criteria', level=1)
    
    doc.add_heading('Entry Criteria', level=2)
    for item in test_plan.get('entry_criteria', []):
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Exit Criteria', level=2)
    for item in test_plan.get('exit_criteria', []):
        doc.add_paragraph(item, style='List Bullet')
    
    # Test Data Requirements
    doc.add_heading('Test Data Requirements', level=1)
    for item in test_plan.get('test_data_requirements', []):
        doc.add_paragraph(item, style='List Bullet')
    
    # Test Environment
    doc.add_heading('Test Environment', level=1)
    env_table = doc.add_table(rows=1, cols=2)
    env_table.style = 'Light Grid Accent 1'
    
    hdr_cells = env_table.rows[0].cells
    hdr_cells[0].text = 'Environment'
    hdr_cells[1].text = 'Purpose'
    
    for env in test_plan.get('test_environment', []):
        row_cells = env_table.add_row().cells
        row_cells[0].text = env.get('name', '')
        row_cells[1].text = env.get('purpose', '')
    
    doc.add_paragraph()
    
    # Testing Activities
    doc.add_heading('Testing Activities', level=1)
    act_table = doc.add_table(rows=1, cols=3)
    act_table.style = 'Light Grid Accent 1'
    
    hdr_cells = act_table.rows[0].cells
    hdr_cells[0].text = 'Activity'
    hdr_cells[1].text = 'Details'
    hdr_cells[2].text = 'Duration'
    
    for activity in test_plan.get('testing_activities', []):
        row_cells = act_table.add_row().cells
        row_cells[0].text = activity.get('activity', '')
        row_cells[1].text = activity.get('details', '')
        row_cells[2].text = activity.get('duration', '')
    
    doc.add_paragraph()
    
    # Roles & Responsibilities
    doc.add_heading('Roles & Responsibilities', level=1)
    role_table = doc.add_table(rows=1, cols=3)
    role_table.style = 'Light Grid Accent 1'
    
    hdr_cells = role_table.rows[0].cells
    hdr_cells[0].text = 'Role'
    hdr_cells[1].text = 'Name'
    hdr_cells[2].text = 'Responsibilities'
    
    for role in test_plan.get('roles_responsibilities', []):
        row_cells = role_table.add_row().cells
        row_cells[0].text = role.get('role', '')
        row_cells[1].text = role.get('name', '')
        row_cells[2].text = role.get('responsibilities', '')
    
    doc.add_paragraph()
    
    # Risks
    doc.add_heading('Risks', level=1)
    for risk in test_plan.get('risks', []):
        doc.add_paragraph(risk, style='List Bullet')
    
    # Assumptions & Dependencies
    doc.add_heading('Assumptions & Dependencies', level=1)
    
    doc.add_heading('Assumptions', level=2)
    for item in test_plan.get('assumptions', []):
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Dependencies', level=2)
    for item in test_plan.get('dependencies', []):
        doc.add_paragraph(item, style='List Bullet')
    
    # Defect Management
    doc.add_heading('Defect Management Process', level=1)
    for item in test_plan.get('defect_management', []):
        doc.add_paragraph(item, style='List Bullet')
    
    # Test Metrics
    doc.add_heading('Test Metrics & KPIs', level=1)
    for metric in test_plan.get('test_metrics', []):
        doc.add_paragraph(metric, style='List Bullet')
    
    # Deliverables
    doc.add_heading('Deliverables', level=1)
    for item in test_plan.get('deliverables', []):
        doc.add_paragraph(item, style='List Bullet')
    
    # Limitations
    doc.add_heading('Limitations & Exclusions', level=1)
    for item in test_plan.get('limitations', []):
        doc.add_paragraph(item, style='List Bullet')
    
    # Approval Section
    doc.add_heading('Approval', level=1)
    approval_table = doc.add_table(rows=1, cols=3)
    approval_table.style = 'Light Grid Accent 1'
    
    hdr_cells = approval_table.rows[0].cells
    hdr_cells[0].text = 'Role'
    hdr_cells[1].text = 'Name'
    hdr_cells[2].text = 'Signature / Date'
    
    approval_roles = [
        ('QA Lead', 'TBD', ''),
        ('Product Manager', 'TBD', '')
    ]
    
    for role_data in approval_roles:
        row_cells = approval_table.add_row().cells
        row_cells[0].text = role_data[0]
        row_cells[1].text = role_data[1]
        row_cells[2].text = role_data[2]
    
    # Save document
    doc.save(output_path)
    print(f"✅ Word document created: {output_path}")


def main():
    """Main function"""
    print("=" * 60)
    print("🚀 Automated Test Plan Generator using Claude API")
    print("=" * 60)
    
    # Check arguments
    if len(sys.argv) < 2:
        print("\n❌ Usage: python generate_test_plan.py <requirements_pdf_path> [output_name]")
        print("\nExample:")
        print("  python generate_test_plan.py requirements.pdf MyProject")
        sys.exit(1)
    
    pdf_path = sys.argv[1]
    project_name = sys.argv[2] if len(sys.argv) > 2 else "Project"
    
    # Check if PDF exists
    if not os.path.exists(pdf_path):
        print(f"❌ Error: PDF file not found: {pdf_path}")
        sys.exit(1)
    
    # Step 1: Read PDF
    requirements_text = read_pdf(pdf_path)
    
    # Step 2: Generate test plan using Claude
    test_plan = generate_test_plan_with_claude(requirements_text, project_name)
    
    # Step 3: Create Word document
    output_docx = f"{project_name}_Test_Plan.docx"
    create_word_document(test_plan, output_docx)
    
    # Save JSON for reference
    json_output = f"{project_name}_Test_Plan.json"
    with open(json_output, 'w', encoding='utf-8') as f:
        json.dump(test_plan, f, indent=2, ensure_ascii=False)
    print(f"✅ JSON saved: {json_output}")
    
    print("\n" + "=" * 60)
    print("✅ Test Plan Generation Complete!")
    print(f"📄 Word Document: {output_docx}")
    print(f"📋 JSON File: {json_output}")
    print("=" * 60)


if __name__ == "__main__":
    main()
